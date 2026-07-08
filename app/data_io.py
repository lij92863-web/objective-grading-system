"""Input/output helpers for teacher files and answer drafts.

All answer sources become the same teacher-review draft. Recognition output is
never treated as final; the web UI must confirm it before grading.
"""

import csv
import json
import re
from pathlib import Path
from typing import Dict, List, Optional


ANSWER_FIELDS = ["question", "answer", "points", "partial_credit", "partial_points", "tags", "difficulty"]
REVIEW_FIELDS = [
    "question",
    "type",
    "answer",
    "points",
    "partial_credit",
    "tags",
    "difficulty",
    "auto_gradable",
    "confidence",
    "warnings",
    "raw_text",
]


def _draft(source_type: str, source_file: Path, items: List[Dict[str, object]], warnings: Optional[List[str]] = None) -> Dict[str, object]:
    return {
        "source_type": source_type,
        "source_file": str(source_file),
        "items": items,
        "warnings": warnings or [],
        "need_teacher_review": True,
    }


def _item(question: object = "", answer: object = "", points: object = 1, raw_text: str = "", **extra: object) -> Dict[str, object]:
    qtype = str(extra.get("type") or "single_choice")
    confidence = extra.get("confidence", 1.0)
    return {
        "question": question,
        "type": qtype,
        "answer": str(answer or "").strip().upper(),
        "points": points or 1,
        "partial_credit": bool(extra.get("partial_credit", qtype == "multiple_choice")),
        "tags": extra.get("tags", ""),
        "difficulty": extra.get("difficulty", ""),
        "auto_gradable": extra.get("auto_gradable", True),
        "confidence": confidence,
        "warnings": extra.get("warnings", []),
        "raw_text": raw_text,
    }


def parse_answer_source(file_path: Path) -> Dict[str, object]:
    suffix = file_path.suffix.lower()
    if suffix == ".csv":
        return parse_csv_answer_source(file_path)
    if suffix in {".xlsx", ".xls"}:
        return parse_excel_answer_source(file_path)
    if suffix == ".docx":
        return parse_docx_answer_source(file_path)
    if suffix == ".pdf":
        return parse_pdf_answer_source(file_path)
    if suffix in {".jpg", ".jpeg", ".png", ".webp"}:
        return parse_image_answer_source(file_path)
    if suffix in {".txt", ".text"}:
        return parse_text_answer_source(file_path)
    return _draft("unknown", file_path, [], ["这个文件类型暂时无法读取，请改用表格导入或手动填写标准答案。"])


def parse_csv_answer_source(file_path: Path) -> Dict[str, object]:
    with file_path.open("r", encoding="utf-8-sig", newline="") as handle:
        reader = csv.DictReader(handle)
        rows = list(reader)
    items = []
    warnings = []
    for index, row in enumerate(rows, start=2):
        question = row.get("question") or row.get("题号") or row.get("q") or row.get("number") or ""
        answer = row.get("answer") or row.get("答案") or row.get("correct") or ""
        if not question:
            warnings.append(f"第 {index} 行缺少题号，请老师补充。")
        if not answer:
            warnings.append(f"第 {index} 行缺少标准答案，请老师补充。")
        qtype = row.get("type") or row.get("题型") or ("multiple_choice" if len(str(answer)) > 1 else "single_choice")
        items.append(
            _item(
                question=question,
                type=qtype,
                answer=answer,
                points=row.get("points") or row.get("分值") or 1,
                partial_credit=str(row.get("partial_credit") or row.get("部分给分") or "").lower() in {"1", "true", "yes", "是"},
                tags=row.get("tags") or row.get("知识点") or "",
                difficulty=row.get("difficulty") or row.get("难度") or "",
                confidence=1.0,
                raw_text=json.dumps(row, ensure_ascii=False),
            )
        )
    return _draft("csv", file_path, items, warnings)


def parse_excel_answer_source(file_path: Path) -> Dict[str, object]:
    try:
        from openpyxl import load_workbook
    except ImportError:
        return _draft("excel", file_path, [], ["这台电脑暂时无法读取 Excel 文件，请另存为 CSV 后再上传，或安装 openpyxl。"])
    workbook = load_workbook(file_path, read_only=True, data_only=True)
    sheet = workbook.active
    rows = list(sheet.iter_rows(values_only=True))
    workbook.close()
    if not rows:
        return _draft("excel", file_path, [], ["这个 Excel 文件是空的，请检查后重新上传。"])
    headers = [str(value or "").strip() for value in rows[0]]
    temp_csv_rows = []
    for values in rows[1:]:
        temp_csv_rows.append({headers[index]: value for index, value in enumerate(values) if index < len(headers)})
    temp = file_path.with_suffix(".converted.csv")
    with temp.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=headers)
        writer.writeheader()
        writer.writerows(temp_csv_rows)
    draft = parse_csv_answer_source(temp)
    draft["source_type"] = "excel"
    draft["source_file"] = str(file_path)
    try:
        temp.unlink()
    except OSError:
        pass
    return draft


def parse_docx_answer_source(file_path: Path) -> Dict[str, object]:
    try:
        from docx import Document
    except ImportError:
        return _draft("docx", file_path, [], ["这个 Word 文件暂时无法读取。请安装 python-docx，或复制答案文字到“手动填写”。"])
    try:
        document = Document(file_path)
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                parts.append(" ".join(cell.text.strip() for cell in row.cells if cell.text.strip()))
        text = "\n".join(parts)
    except Exception:
        return _draft("docx", file_path, [], ["这个 Word 文件暂时无法读取，请确认文件没有损坏，或改用手动填写标准答案。"])
    draft = parse_text_answer_source(file_path, text=text)
    draft["source_type"] = "docx"
    draft["warnings"].append("已从 Word 中提取文字并生成草稿，请老师逐题确认。")
    return draft


def parse_pdf_answer_source(file_path: Path) -> Dict[str, object]:
    text = ""
    try:
        try:
            from pypdf import PdfReader
        except ImportError:
            from PyPDF2 import PdfReader  # type: ignore
        reader = PdfReader(str(file_path))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception:
        return _draft("pdf", file_path, [], ["这个 PDF 暂时无法读取。若是扫描版，请使用图片/拍照方式导入。"])
    if not text.strip():
        return _draft("scanned_pdf", file_path, [], ["这个 PDF 像是扫描版，没有可直接提取的文字。请使用图片识别或手动填写，识别后仍需老师确认。"])
    draft = parse_text_answer_source(file_path, text=text)
    draft["source_type"] = "pdf"
    draft["warnings"].append("已从 PDF 中提取文字并生成草稿，请老师逐题确认。")
    return draft


def parse_image_answer_source(file_path: Path) -> Dict[str, object]:
    return _draft(
        "image",
        file_path,
        [],
        ["图片已保存。当前本地版不会直接把图片结果作为正式答案，请在确认表中手动补充，或后续接入千问视觉生成草稿。"],
    )


def parse_text_answer_source(file_path: Path, text: Optional[str] = None) -> Dict[str, object]:
    if text is None:
        text = file_path.read_text(encoding="utf-8-sig")
    items = []
    warnings = []
    pattern = re.compile(r"^\s*(\d+)\s*[\.\、:：\)]\s*([A-Ha-h0-9|/;；,，]+)", re.MULTILINE)
    for match in pattern.finditer(text):
        answer = match.group(2).replace("，", "").replace(",", "").replace("；", "|").replace(";", "|")
        qtype = "multiple_choice" if answer.isalpha() and len(answer) > 1 else "single_choice"
        items.append(_item(question=match.group(1), type=qtype, answer=answer, points=1, confidence=0.78, raw_text=match.group(0)))
    if not items:
        warnings.append("没有稳定识别出题号和答案，请老师使用手动填写。")
    else:
        warnings.append("已按本地规则生成标准答案草稿，请老师检查黄色标记内容。")
    return _draft("text", file_path, items, warnings)


def draft_to_review_rows(draft: Dict[str, object]) -> List[Dict[str, object]]:
    rows = []
    for item in draft.get("items", []):
        if not isinstance(item, dict):
            continue
        rows.append({field: item.get(field, "") for field in REVIEW_FIELDS})
    return rows


def review_rows_to_answer_key_csv(rows: List[Dict[str, object]], out_path: Path) -> None:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    with out_path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=ANSWER_FIELDS)
        writer.writeheader()
        for row in rows:
            writer.writerow(
                {
                    "question": row.get("question", ""),
                    "answer": row.get("answer", ""),
                    "points": row.get("points", 1),
                    "partial_credit": row.get("partial_credit", "false"),
                    "partial_points": row.get("partial_points", ""),
                    "tags": row.get("tags", ""),
                    "difficulty": row.get("difficulty", ""),
                }
            )

